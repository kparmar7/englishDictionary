import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
import os

cwd = os.getcwd()
print(cwd) #D:\SCRAPY\FIRST\englishDictionary

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(51,51,51)

with open("itemsA.json") as j:
    j_data = json.load(j)
    #print(len(j_data))

    #for j in range(len(j_data)):
        #DATA = j_data[j]["data"]
        #print(DATA)
        #print(len(DATA))
        #for d in range(len(DATA)):
            #print(DATA[d])

    DATA = j_data[0]["data"][0]
    HEAD = DATA[0]
    head_data = HEAD['Head']
    #print(head_data)
    #try:
    #    searchbox_result = re.match("^.*(?=(\())", searchbox.group()
    #except:
    #    searchbox_result = None
    classs = re.search('class="(.*?)"', head_data).group(1)
    #print(classs)
    if classs == 'Head':
        value = re.search('>(.*?)<', head_data).group(1)
        para = document.add_paragraph(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1) #<class="head_data">
        head_data = head_data.replace(value, '', 1)          #remove value
        #print(head_data)
        classs = re.search('class="(.*?)"', head_data).group(1) # reset search classs
    if classs == "HWD":
        value = re.search('>(.*?)<', head_data).group(1)
        style = para.add_run(value)
        style.font.size = Pt(17)
        style.font.bold = True
        para = document.add_paragraph()
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        #print(head_data)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "HYPHENATION":
        value = re.search('>(.*?)<', head_data).group(1)
        style = para.add_run(value)
        style.font.size = Pt(17)
        style.font.color.rgb = RGBColor(255,0,0)
        style.font.bold = True

        
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        #print(head_data)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "HOMNUM":
        value = re.search('>(.*?)<', head_data).group(1)
        style = para.add_run(value)
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(255,0,0)
        style.font.bold = True
        
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        #print(head_data)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "Variant":
        value = re.search('>(.*?)<', head_data).group(1)
        style = para.add_run(value)
        style.font.size = Pt(17)
        style.font.color.rgb = RGBColor(255,0,0)
        style.font.bold = True
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "neutral span":
        value = re.search('>(.*?)<', head_data).group(1)
        style = para.add_run(value)
        style.font.bold = True
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "ORTHVAR":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "PronCodes":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        #print(head_data)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(classs)
    if classs == "neutral span":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "PRON":
        value = re.search('>(.*?)<', head_data).group(1)
        style = para.add_run(value)
        style.font.color.rgb = RGBColor(0,128,0)
        style.font.bold = True
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "neutral span":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "POS":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "Inflections":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "neutral span":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "PLURALFORM":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        #print(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "italic span":
        value = re.search('>(.*?)<', head_data).group(1)
        li = value.split(' ',1)
        style = para.add_run(li[0] + ' ')
        style.font.italic = True
        style.font.name = 'Times New Roman'
        style = para.add_run(li[1])
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "PLURALFORM":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
        #print(head_data)
    if classs == "neutral span":
        
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        classs = re.search('class="(.*?)"', head_data).group(1)
    if classs == "neutral span":
        value = re.search('>(.*?)<', head_data).group(1)
        para.add_run(value)
        head_data = re.sub('<class="(.*?)>','',head_data, 1)
        head_data = head_data.replace(value, '', 1)
        
        try:
            classs = re.search('class="(.*?)"', head_data).group(1)
        except:
            head_data = head_data.replace(head_data, '')
        #print(head_data)
    document.save("my_written_file.docx")
    Sense = DATA[1]['Sense']
    #print(Sense[0]["1"])

